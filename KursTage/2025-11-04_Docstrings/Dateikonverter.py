from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
#Nicht vergessen, die benötigten Pakete zu installieren:
#pip install python-docx reportlab

class Dateikonverter:
    '''
    Dateikonverter ist eine Klasse die mehrere Datei Konvertierungsmoeglichkeiten
    bereitstellt. Es kann zu pdf, md oder word Konvertieren.
    '''

    def __init__():
        '''
        Constructor to initialize a Dateikonverter Object

        Args:
            input_file: Pfad zu eine Datei die Konvertiert werden soll.
            Erlaubte Datei Typen sind: .txt, .md, .py
            Should work for all text based files.
        '''
        

    def _read_text(input_file):
        '''
        Liest die eingegebene Datei in self.input_file ein.

        Return:
            Die inhalte der geoeffneten Datei.
        '''
        with open(input_file, "r", encoding="utf-8") as f:
            return f.read()

    def to_markdown(input_file, output_file):
        '''

        '''
        text = Dateikonverter._read_text(input_file)
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(f"# Converted File\n\n{text}")

    def to_word(input_file, output_file):
        text = Dateikonverter._read_text(input_file)
        doc = Document()
        doc.add_heading("Converted File", level=1)
        doc.add_paragraph(text)
        doc.save(output_file)

    def to_pdf(input_file, output_file):
        text = Dateikonverter._read_text(input_file)
        c = canvas.Canvas(output_file, pagesize=A4)
        width, height = A4
        y = height - 50
        for line in text.splitlines():
            c.drawString(50, y, line)
            y -= 15
            if y < 50:
                c.showPage()
                y = height - 50
        c.save()

#Beispielnutzung:
# test_konvertierer = Dateikonverter("Data/konvertierteDatei.md")
# test_konvertierer.to_markdown("Data/konvertierteDatei2.txt")

Dateikonverter.to_pdf('Data/konvertierteDatei2.txt','Data/test1.pdf')