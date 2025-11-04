from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

# Nicht vergessen, die benötigten Pakete zu installieren:
# pip install python-docx reportlab


class Dateikonverter:
    def __init__(self, input_file):
        self.input_file = input_file

    def _read_text(self):
        with open(self.input_file, "r", encoding="utf-8") as f:
            return f.read()

    def to_markdown(self, output_file):
        text = self._read_text()
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(f"# Converted File\n\n{text}")

    def to_word(self, output_file):
        text = self._read_text()
        doc = Document()
        doc.add_heading("Converted File", level=1)
        doc.add_paragraph(text)
        doc.save(output_file)

    def to_pdf(self, output_file):
        text = self._read_text()
        c = canvas.Canvas(output_file, pagesize=A4)
        width, height = A4
        y = height - 50
        for line in text.splitlines():
            c.drawString(50, y, line)
            y -= 15
            if y < 50:
                c.showPage()
                y = height - 50
        c.save()


test_konvertierer = Dateikonverter("converter_programm/zuKonvertierendeDatei.txt")
test_konvertierer.to_pdf("converter_programm/konvertierteDatei.pdf")
