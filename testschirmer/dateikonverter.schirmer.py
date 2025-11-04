from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
#Nicht vergessen, die benötigten Pakete zu installieren:
#pip install python-docx reportlab

class Dateikonverter: 
    """Hier wird die klasse definiert die alle funktionen als die class Dateikonverter zusammenfügt die ab zeile 8 bis zeile 38 zusammenfügt"""
    def __init__(self, input_file):
        self.input_file = input_file
    """Hier wird ein eingabe datei defininiert die in der klasse benutzt wird"""
    def _read_text(self):
        with open(self.input_file, "r", encoding="utf-8") as f:
            return f.read()
    """Hier wird die datei gelesen"""
    def to_markdown(self, output_file):
        text = self._read_text()
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(f"# Converted File\n\n{text}")
    """Hier wird die datei in markdown konvertiert"""
    def to_word(self, output_file):
        text = self._read_text()
        doc = Document()
        doc.add_heading("Converted File", level=1)
        doc.add_paragraph(text)
        doc.save(output_file)
    """Hier wird die datei in word konvertiert"""
    def to_pdf(self, output_file):
        text = self._read_text()
        c = canvas.Canvas(output_file, pagesize=A4)
        width, height = A4
        y = height - 50
        for line in text.splitlines():
            c.drawString(50, y, line)
            y -= 15
            if y < 50:
                c.showPage()
                y = height - 50
        c.save()
    """Hier wird die datei in pdf konvertiert"""
test_konvertierer = Dateikonverter("zuKonvertierendeDatei.txt") #Hier wird die datei definiert die konvertiert werden soll
test_konvertierer.to_pdf("konvertierteDatei.pdf")#Hier wird die datei in pdf konvertiert