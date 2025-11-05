from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
"""Importiert die nözigen Pakete"""
#Nicht vergessen, die benötigten Pakete zu installieren:
#pip install python-docx reportlab

"""
Eine Klasse die fürs konvertieren in die verschiedenen Textformate da ist.
"""
class Dateikonverter:
    def __init__(self, input_file):
        self.input_file = input_file

    """
    Das liest den inhalt der datei und gibt ein string zurück
    """
    def _read_text(self):
        with open(self.input_file, "r", encoding="utf-8") as f:
            return f.read()

    """ Erstellt mit der Datei eine Markdown Datei? Und fügt converted file mit hinzu.
    """
    def to_markdown(self, output_file):
        text = self._read_text()
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(f"# Converted File\n\n{text}")

    """ Konvertiert die Datei in eine .docx Datei.
    """
    def to_word(self, output_file):
        text = self._read_text()
        doc = Document()
        doc.add_heading("Converted File", level=1)
        doc.add_paragraph(text)
        doc.save(output_file)

    """ Hier wird die .docx in .pdf konvertiert. Im A4 Format. Und dann irschendwo abgespeichert
    """
    def to_pdf(self, output_file):
        text = self._read_text()
        c = canvas.Canvas(output_file, pagesize=A4)
        width, height = A4
        y = height - 50
        for line in text.splitlines():
            c.drawString(50, y, line)
            y -= 15
            if y < 50:
                c.showPage()
                y = height - 50
        c.save()

""" Konvertierende datei test dingens
"""
test_konvertierer = Dateikonverter("zuKonvertierendeDatei.txt")
test_konvertierer.to_pdf("konvertierteDatei.pdf")